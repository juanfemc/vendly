from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "politicas"
DOCX_PATH = OUT / "politica-de-seguridad-vendly.docx"

BLACK = "111111"
INK = "1D2433"
MUTED = "667085"
ORANGE = "FF6A00"
SOFT_ORANGE = "FFF3EA"
SOFT_GRAY = "F4F6F8"
BORDER = "E4E7EC"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_twips=9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")


def set_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def para(doc, text="", size=11, color=INK, bold=False, italic=False, before=0, after=6, align=None, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def heading(doc, text, level=1):
    sizes = {1: 16, 2: 13, 3: 12}
    before = {1: 16, 2: 11, 3: 8}
    after = {1: 7, 2: 5, 3: 4}
    color = ORANGE if level == 1 else BLACK
    p = para(doc, text, size=sizes[level], color=color, bold=True, before=before[level], after=after[level])
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    set_font(run, size=10.5, color=INK)
    return p


def callout(doc, title, body, fill=SOFT_ORANGE):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "FFD2B3")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_font(r, 10.5, BLACK, True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(3)
    r2 = p2.add_run(body)
    set_font(r2, 10, INK)
    para(doc, "", after=3)


def label_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table)
    widths = [Inches(1.75), Inches(4.75)]
    for row_idx, (label, value) in enumerate(rows):
        for col_idx, text in enumerate((label, value)):
            cell = table.cell(row_idx, col_idx)
            cell.width = widths[col_idx]
            set_cell_border(cell)
            set_cell_shading(cell, SOFT_GRAY if col_idx == 0 else WHITE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            set_font(r, 10, INK, bold=(col_idx == 0))
    para(doc, "", after=4)


def add_header_footer(doc):
    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0]
        hp.text = ""
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = hp.add_run("VendlySuite | Politica de Seguridad")
        set_font(r, 8.5, MUTED)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run("Documento de referencia - completar datos legales antes de publicar")
        set_font(r, 8.5, MUTED)


def cover(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.72)
    sec.right_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.85)

    para(doc, "VENDLYSUITE", size=11, color=ORANGE, bold=True, after=44, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Politica de Seguridad", size=30, color=BLACK, bold=True, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(
        doc,
        "Proteccion de la informacion, cuentas, tiendas, clientes, WhatsApp, pagos e inteligencia artificial.",
        size=12.5,
        color=MUTED,
        after=34,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    table = doc.add_table(rows=1, cols=3)
    set_table_width(table)
    items = [
        ("Confidencialidad", "Acceso solo para personas y sistemas autorizados."),
        ("Integridad", "Datos protegidos contra alteraciones no autorizadas."),
        ("Disponibilidad", "Continuidad razonable del servicio y recuperacion."),
    ]
    for i, (title, text) in enumerate(items):
        cell = table.cell(0, i)
        set_cell_shading(cell, SOFT_GRAY)
        set_cell_border(cell, BORDER)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        set_font(r, 11, BLACK, True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(text)
        set_font(r2, 9.2, MUTED)

    para(doc, "", after=38)
    callout(
        doc,
        "Nota legal",
        "Este documento es una base de trabajo. Antes de publicarlo debe completarse con razon social, NIT, ciudad, correos oficiales y revision juridica final.",
    )
    label_table(
        doc,
        [
            ("Version", "1.0"),
            ("Ultima actualizacion", "[Dia] de [Mes] de [Ano]"),
            ("Responsable", "[Razon social / VendlySuite]"),
            ("Contacto de seguridad", "[correo de seguridad]"),
        ],
    )
    doc.add_section(WD_SECTION.NEW_PAGE)


def build_doc():
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    cover(doc)

    sec = doc.sections[-1]
    sec.top_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.78)
    sec.left_margin = Inches(0.85)

    add_header_footer(doc)

    para(doc, "Resumen ejecutivo", size=18, color=BLACK, bold=True, after=8)
    para(
        doc,
        "Esta politica define los controles, responsabilidades y buenas practicas que VendlySuite adopta para proteger la informacion de usuarios, tiendas, clientes finales y servicios integrados.",
        size=11.3,
        color=INK,
        after=8,
    )
    callout(
        doc,
        "Objetivo",
        "Preservar la confidencialidad, integridad y disponibilidad de la informacion tratada en la plataforma, reduciendo riesgos operativos, legales y de ciberseguridad.",
    )

    sections = [
        (
            "1. Alcance",
            [
                "Esta politica aplica a usuarios que crean una cuenta en Vendly, propietarios o administradores de tiendas, clientes finales que compran o interactuan con tiendas creadas en la plataforma, visitantes del sitio web, personal autorizado, proveedores y terceros que puedan tener acceso a informacion del sistema.",
                "Tambien aplica a la informacion tratada en el sitio web, panel administrativo, tiendas publicadas, formularios, integraciones, sistemas de mensajeria, servicios de pago, herramientas de inteligencia artificial y demas funcionalidades disponibles.",
            ],
        ),
        (
            "2. Principios de seguridad",
            [
                "Vendly protege la informacion bajo principios de confidencialidad, integridad, disponibilidad, minimizacion, responsabilidad y trazabilidad.",
                "Solo se deben solicitar y tratar los datos necesarios para prestar el servicio, operar la plataforma, prevenir abuso, cumplir obligaciones legales y mejorar la seguridad.",
            ],
        ),
        (
            "3. Informacion que protegemos",
            [
                "Vendly puede tratar nombres, correos, telefonos, contrasenas, informacion de tiendas, productos, categorias, imagenes, precios, pedidos, direcciones de envio, ciudades, departamentos, metodos de entrega, configuraciones de dominios, redes sociales, reseñas, logs tecnicos y eventos de seguridad.",
                "La plataforma no debe solicitar contrasenas bancarias, codigos OTP financieros, claves privadas, datos sensibles innecesarios ni informacion que no sea requerida para el funcionamiento del servicio.",
            ],
        ),
        (
            "4. Proteccion de cuentas y contrasenas",
            [
                "Los usuarios son responsables de mantener la confidencialidad de sus credenciales de acceso y de cerrar sesion en dispositivos compartidos.",
                "Vendly recomienda usar contrasenas seguras y unicas, mantener actualizado el correo de recuperacion y reportar cualquier actividad sospechosa.",
                "La plataforma podra bloquear, limitar o suspender cuentas cuando detecte actividad sospechosa, abuso, fraude o riesgos para otros usuarios.",
            ],
        ),
        (
            "5. Control de acceso",
            [
                "Vendly implementa controles para que cada usuario acceda unicamente a la informacion y funcionalidades correspondientes a su rol, plan o tienda.",
                "Esto incluye separacion entre administrador general, propietario de tienda y cliente final; restricciones por plan; validaciones de permisos; proteccion de rutas administrativas; y registro de eventos relevantes.",
            ],
        ),
        (
            "6. Seguridad en comunicaciones",
            [
                "Vendly debe operar mediante conexiones seguras HTTPS/TLS en produccion.",
                "Las comunicaciones con pasarelas de pago, WhatsApp Business API, proveedores de correo, almacenamiento, inteligencia artificial, verificacion antiabuso y demas servicios externos deben realizarse mediante canales seguros y credenciales protegidas.",
            ],
        ),
        (
            "7. Pagos y metodos de pago",
            [
                "Vendly puede permitir que las tiendas configuren metodos de pago como WhatsApp, pagos manuales, transferencias, pasarelas externas o proveedores como Mercado Pago, segun el plan y la configuracion disponible.",
                "Cuando se usen pasarelas externas, la informacion de pago sera procesada por el proveedor correspondiente. Vendly no debe almacenar datos completos de tarjetas, codigos CVV ni credenciales bancarias.",
            ],
        ),
        (
            "8. Seguridad en WhatsApp Business API",
            [
                "Cuando Vendly utilice WhatsApp Business API para enviar mensajes de verificacion, bienvenida, pedidos o notificaciones, se aplicaran controles como uso de tokens protegidos, plantillas aprobadas, consentimiento cuando aplique y registro limitado de eventos necesarios.",
                "Los mensajes enviados por WhatsApp dependen de disponibilidad, reglas, aprobacion de plantillas y politicas de Meta.",
            ],
        ),
        (
            "9. Inteligencia artificial",
            [
                "Vendly puede ofrecer funciones de inteligencia artificial para generar o mejorar nombres de productos, descripciones, caracteristicas, etiquetas, avisos promocionales, imagenes o portadas.",
                "El usuario debe evitar ingresar informacion sensible, confidencial o de terceros sin autorizacion en las herramientas de IA. Los resultados generados por IA deben ser revisados antes de publicarse.",
            ],
        ),
        (
            "10. Proteccion contra abuso y fraude",
            [
                "Vendly puede usar mecanismos como verificacion por WhatsApp, Cloudflare Turnstile, limites de intentos, bloqueo temporal, validacion de numeros, correos o dominios y registro de actividad tecnica.",
                "Vendly podra negar, suspender o limitar el acceso cuando existan indicios razonables de abuso, manipulacion, uso fraudulento o afectacion a la seguridad del sistema.",
            ],
        ),
        (
            "11. Almacenamiento y copias de seguridad",
            [
                "Vendly puede realizar copias de seguridad, respaldos tecnicos y registros operativos para mantener continuidad del servicio, prevenir perdida de informacion y facilitar recuperacion ante incidentes.",
                "La restauracion de informacion estara sujeta a disponibilidad tecnica, alcance del incidente y condiciones del servicio contratado.",
            ],
        ),
        (
            "12. Proveedores y servicios de terceros",
            [
                "Vendly puede apoyarse en proveedores de hosting, dominios, DNS, pasarelas de pago, WhatsApp Business API, correo, almacenamiento, analitica, inteligencia artificial, proteccion antiabuso, seguridad y monitoreo.",
                "Cada proveedor puede tener sus propios terminos, politicas, disponibilidad y responsabilidades.",
            ],
        ),
        (
            "13. Reporte de vulnerabilidades",
            [
                "Cualquier persona que identifique una posible vulnerabilidad debe reportarla de forma responsable al correo [correo de seguridad o soporte].",
                "No esta permitido explotar vulnerabilidades, acceder a informacion de terceros, modificar datos, interrumpir servicios, realizar ataques de denegacion de servicio, ingenieria social, spam, extraccion masiva de datos o cualquier actividad ilegal.",
            ],
        ),
        (
            "14. Gestion de incidentes de seguridad",
            [
                "Ante un incidente, Vendly podra investigar, contener accesos no autorizados, revocar tokens, forzar cambios de contrasena, desactivar funciones afectadas, restaurar servicios, notificar a usuarios afectados e informar a autoridades cuando corresponda.",
                "La respuesta dependera de la naturaleza, alcance y severidad del incidente.",
            ],
        ),
        (
            "15. Responsabilidades del usuario",
            [
                "El usuario se compromete a usar la plataforma de forma legal, no acceder a cuentas de terceros, no cargar contenido malicioso o ilegal, no usar Vendly para phishing, spam, estafas o suplantacion y cumplir las normas de proteccion de datos respecto a sus propios clientes.",
                "Cada tienda es responsable de la informacion que publica, los productos que vende, las condiciones comerciales que ofrece y el tratamiento de los datos de sus clientes finales.",
            ],
        ),
        (
            "16. Tratamiento de datos personales",
            [
                "Vendly trata datos personales conforme a la normativa aplicable de proteccion de datos personales, incluyendo la Ley 1581 de 2012 en Colombia, cuando corresponda.",
                "Los titulares podran ejercer sus derechos de conocer, actualizar, rectificar, suprimir informacion y revocar autorizaciones de acuerdo con la ley aplicable y la politica de tratamiento de datos personales de Vendly.",
            ],
        ),
        (
            "17. Retencion de informacion",
            [
                "Vendly conservara la informacion durante el tiempo necesario para prestar el servicio, cumplir obligaciones legales, resolver disputas, prevenir fraude, mejorar la seguridad y atender requerimientos administrativos, contables o judiciales.",
                "Cuando la informacion deje de ser necesaria, Vendly podra eliminarla, anonimizarla o conservarla bloqueada cuando exista obligacion legal o interes legitimo de seguridad.",
            ],
        ),
        (
            "18. Cambios en esta politica",
            [
                "Vendly podra actualizar esta Politica de Seguridad para reflejar cambios legales, tecnicos, operativos o comerciales.",
                "Cuando los cambios sean relevantes, Vendly podra informar a los usuarios mediante la pagina web, correo electronico, panel administrativo u otros canales disponibles.",
            ],
        ),
        (
            "19. Contacto",
            [
                "Nombre comercial: VendlySuite. Correo de soporte: [correo]. Correo de seguridad: [correo]. WhatsApp de soporte: [numero]. Sitio web: [dominio]. Ciudad y pais: [ciudad], Colombia.",
            ],
        ),
    ]

    for title, paragraphs in sections:
        heading(doc, title, 1)
        for text in paragraphs:
            para(doc, text, after=6)

    heading(doc, "Fuentes de referencia", 1)
    for item in [
        "Ley 1581 de 2012 - Regimen General de Proteccion de Datos Personales en Colombia.",
        "Superintendencia de Industria y Comercio - lineamientos de proteccion de datos personales.",
        "Politicas y documentacion aplicable de proveedores externos como Meta, pasarelas de pago, Cloudflare y servicios de inteligencia artificial.",
    ]:
        bullet(doc, item)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
